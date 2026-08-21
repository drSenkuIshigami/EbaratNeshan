"""Turn recovered Markdown into HTML, DOCX, or PDF on this computer."""

from __future__ import annotations

import html
import os
import re
import shutil
import subprocess
import tempfile
from io import BytesIO
from pathlib import Path

from .postprocess import has_persian

HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
UL = re.compile(r"^(\s*)[-*+]\s+(.*)$")
OL = re.compile(r"^(\s*)(\d+)[.)]\s+(.*)$")
FENCE = re.compile(r"^```")
IMG = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
BOLD = re.compile(r"\*\*(.+?)\*\*")
ITALIC = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")


def export_bytes(
    markdown: str,
    fmt: str,
    title: str = "document",
    base_dir: Path | None = None,
) -> tuple[bytes, str, str]:
    """Return (payload, filename, content_type)."""
    fmt = fmt.lower().strip()
    stem = _stem(title)
    rtl = has_persian(markdown)
    body = _strip_yaml(markdown)
    if base_dir:
        body = _absolutize_images(body, Path(base_dir))
    if fmt == "html":
        data = markdown_to_html(body, title=stem, rtl=rtl).encode("utf-8")
        return data, f"{stem}.html", "text/html; charset=utf-8"
    if fmt == "txt":
        data = body.encode("utf-8")
        return data, f"{stem}.txt", "text/plain; charset=utf-8"
    if fmt in {"docx", "doc"}:
        data = markdown_to_docx(body, title=stem, rtl=rtl)
        return data, f"{stem}.docx", (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    if fmt == "pdf":
        html_text = markdown_to_html(body, title=stem, rtl=rtl, for_print=True)
        pdf = html_to_pdf(html_text)
        if pdf:
            return pdf, f"{stem}.pdf", "application/pdf"
        data = html_text.encode("utf-8")
        return data, f"{stem}.print.html", "text/html; charset=utf-8"
    raise ValueError("Use html, docx, pdf, or txt.")


def markdown_to_html(markdown: str, *, title: str, rtl: bool, for_print: bool = False) -> str:
    lang = "fa" if rtl else "en"
    direction = "rtl" if rtl else "ltr"
    inner = _blocks_to_html(_parse_blocks(markdown))
    print_js = (
        "<script>window.addEventListener('load',function(){setTimeout(function(){window.print()},200)});</script>"
        if for_print
        else ""
    )
    return f"""<!DOCTYPE html>
<html lang="{lang}" dir="{direction}">
<head>
<meta charset="utf-8">
<title>{html.escape(title)}</title>
<style>
  body {{ font-family: "Segoe UI", Tahoma, "Noto Naskh Arabic", Arial, sans-serif;
         line-height: 1.65; max-width: 800px; margin: 2rem auto; padding: 0 1.2rem;
         color: #111; }}
  pre, code {{ font-family: Consolas, "Courier New", monospace; direction: ltr; text-align: left; }}
  pre {{ background: #f4f4f4; padding: 0.8rem; overflow: auto; }}
  table {{ border-collapse: collapse; width: 100%; margin: 1rem 0; }}
  th, td {{ border: 1px solid #bbb; padding: 0.4rem 0.55rem; }}
  img {{ max-width: 100%; }}
  @media print {{ body {{ margin: 0; max-width: none; }} }}
</style>
</head>
<body>
{inner}
{print_js}
</body>
</html>
"""


def markdown_to_docx(markdown: str, *, title: str, rtl: bool) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    if rtl:
        doc.sections[0]._sectPr.append(OxmlElement("w:bidi"))

    def _bidi(paragraph) -> None:
        if not rtl:
            return
        pPr = paragraph._p.get_or_add_pPr()
        if pPr.find(qn("w:bidi")) is None:
            pPr.append(OxmlElement("w:bidi"))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for block in _parse_blocks(markdown):
        kind = block[0]
        if kind == "h":
            _, level, text = block
            p = doc.add_heading(_plain(text), level=min(max(level, 1), 9))
            _bidi(p)
        elif kind == "p":
            p = doc.add_paragraph(_plain(block[1]))
            _bidi(p)
        elif kind == "pre":
            p = doc.add_paragraph(block[1])
            _bidi(p)
        elif kind == "ul":
            for item in block[1]:
                p = doc.add_paragraph(_plain(item), style="List Bullet")
                _bidi(p)
        elif kind == "ol":
            for item in block[1]:
                p = doc.add_paragraph(_plain(item), style="List Number")
                _bidi(p)
        elif kind == "table":
            rows = block[1]
            if not rows:
                continue
            table = doc.add_table(rows=len(rows), cols=max(len(rows[0]), 1))
            table.style = "Table Grid"
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row):
                    if ci < len(table.rows[ri].cells):
                        table.rows[ri].cells[ci].text = _plain(cell)
        elif kind == "img":
            _, alt, src = block
            path = Path(src)
            if path.is_file():
                try:
                    doc.add_picture(str(path))
                except Exception:
                    p = doc.add_paragraph(alt or src)
                    _bidi(p)
            else:
                p = doc.add_paragraph(alt or src)
                _bidi(p)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def html_to_pdf(html_text: str) -> bytes | None:
    chrome = _chrome_path()
    if not chrome:
        return None
    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "doc.html"
        dest = Path(tmp) / "doc.pdf"
        src.write_text(html_text, encoding="utf-8")
        try:
            subprocess.run(
                [
                    chrome,
                    "--headless=new",
                    "--disable-gpu",
                    "--no-pdf-header-footer",
                    f"--print-to-pdf={dest}",
                    str(src),
                ],
                check=True,
                capture_output=True,
                timeout=60,
            )
        except (OSError, subprocess.SubprocessError, subprocess.TimeoutExpired):
            return None
        if dest.is_file() and dest.stat().st_size > 0:
            return dest.read_bytes()
    return None


def _chrome_path() -> str | None:
    env = os.environ.get("CHROME_PATH") or os.environ.get("EDGE_PATH")
    if env and Path(env).is_file():
        return env
    for name in ("msedge", "chrome", "chromium", "chromium-browser", "google-chrome"):
        found = shutil.which(name)
        if found:
            return found
    win_paths = [
        Path(os.environ.get("PROGRAMFILES", r"C:\Program Files"))
        / "Microsoft/Edge/Application/msedge.exe",
        Path(os.environ.get("PROGRAMFILES(X86)", r"C:\Program Files (x86)"))
        / "Microsoft/Edge/Application/msedge.exe",
        Path(os.environ.get("PROGRAMFILES", r"C:\Program Files"))
        / "Google/Chrome/Application/chrome.exe",
        Path(os.environ.get("LOCALAPPDATA", "")) / "Google/Chrome/Application/chrome.exe",
    ]
    for path in win_paths:
        if path.is_file():
            return str(path)
    return None


def _absolutize_images(markdown: str, base: Path) -> str:
    def repl(match: re.Match) -> str:
        src = match.group(2)
        path = Path(src)
        if not path.is_absolute():
            path = (base / src).resolve()
        return f"![{match.group(1)}]({path.as_posix()})"

    return IMG.sub(repl, markdown)


def _parse_blocks(markdown: str) -> list[tuple]:
    lines = markdown.replace("\r\n", "\n").split("\n")
    blocks: list[tuple] = []
    i = 0
    para: list[str] = []

    def flush_para() -> None:
        text = " ".join(p.strip() for p in para if p.strip()).strip()
        para.clear()
        if text:
            blocks.append(("p", text))

    while i < len(lines):
        line = lines[i]
        if FENCE.match(line.strip()):
            flush_para()
            i += 1
            chunk: list[str] = []
            while i < len(lines) and not FENCE.match(lines[i].strip()):
                chunk.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            blocks.append(("pre", "\n".join(chunk)))
            continue
        if _is_table_row(line) and i + 1 < len(lines) and _is_table_sep(lines[i + 1]):
            flush_para()
            rows = [_split_row(line)]
            i += 2
            while i < len(lines) and _is_table_row(lines[i]):
                rows.append(_split_row(lines[i]))
                i += 1
            blocks.append(("table", rows))
            continue
        hm = HEADING.match(line)
        if hm:
            flush_para()
            blocks.append(("h", len(hm.group(1)), hm.group(2).strip()))
            i += 1
            continue
        um = UL.match(line)
        if um:
            flush_para()
            items = [um.group(2)]
            i += 1
            while i < len(lines) and UL.match(lines[i]):
                items.append(UL.match(lines[i]).group(2))
                i += 1
            blocks.append(("ul", items))
            continue
        om = OL.match(line)
        if om:
            flush_para()
            items = [om.group(3)]
            i += 1
            while i < len(lines) and OL.match(lines[i]):
                items.append(OL.match(lines[i]).group(3))
                i += 1
            blocks.append(("ol", items))
            continue
        im = IMG.search(line.strip())
        if im and line.strip().startswith("!"):
            flush_para()
            blocks.append(("img", im.group(1), im.group(2)))
            i += 1
            continue
        if not line.strip():
            flush_para()
            i += 1
            continue
        para.append(line)
        i += 1
    flush_para()
    return blocks


def _blocks_to_html(blocks: list[tuple]) -> str:
    parts: list[str] = []
    for block in blocks:
        kind = block[0]
        if kind == "h":
            _, level, text = block
            parts.append(f"<h{min(level, 6)}>{_inline_html(text)}</h{min(level, 6)}>")
        elif kind == "p":
            parts.append(f"<p>{_inline_html(block[1])}</p>")
        elif kind == "pre":
            parts.append(f"<pre><code>{html.escape(block[1])}</code></pre>")
        elif kind == "ul":
            items = "".join(f"<li>{_inline_html(x)}</li>" for x in block[1])
            parts.append(f"<ul>{items}</ul>")
        elif kind == "ol":
            items = "".join(f"<li>{_inline_html(x)}</li>" for x in block[1])
            parts.append(f"<ol>{items}</ol>")
        elif kind == "table":
            rows = block[1]
            if not rows:
                continue
            head = "".join(f"<th>{_inline_html(c)}</th>" for c in rows[0])
            body = []
            for row in rows[1:]:
                body.append("<tr>" + "".join(f"<td>{_inline_html(c)}</td>" for c in row) + "</tr>")
            parts.append(
                f"<table><thead><tr>{head}</tr></thead><tbody>{''.join(body)}</tbody></table>"
            )
        elif kind == "img":
            _, alt, src = block
            parts.append(
                f'<p><img src="{html.escape(src, quote=True)}" alt="{html.escape(alt)}"></p>'
            )
    return "\n".join(parts)


def _inline_html(text: str) -> str:
    def repl_img(m: re.Match) -> str:
        return f'<img src="{html.escape(m.group(2), quote=True)}" alt="{html.escape(m.group(1))}">'

    def repl_link(m: re.Match) -> str:
        return f'<a href="{html.escape(m.group(2), quote=True)}">{html.escape(m.group(1))}</a>'

    text = IMG.sub(repl_img, text)
    text = LINK.sub(repl_link, text)
    pieces = re.split(r"(<[^>]+>)", text)
    out = []
    for piece in pieces:
        if piece.startswith("<"):
            out.append(piece)
            continue
        esc = html.escape(piece)
        esc = BOLD.sub(r"<strong>\1</strong>", esc)
        esc = ITALIC.sub(r"<em>\1</em>", esc)
        out.append(esc)
    return "".join(out)


def _plain(text: str) -> str:
    text = IMG.sub(r"\1", text)
    text = LINK.sub(r"\1", text)
    text = BOLD.sub(r"\1", text)
    text = ITALIC.sub(r"\1", text)
    return text.replace("`", "")


def _strip_yaml(markdown: str) -> str:
    if markdown.startswith("---"):
        end = markdown.find("\n---", 3)
        if end != -1:
            return markdown[end + 4 :].lstrip("\n")
    return markdown


def _is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and s.count("|") >= 2


def _is_table_sep(line: str) -> bool:
    return _is_table_row(line) and not (
        line.strip().strip("|").replace(":", "").replace("-", "").replace("|", "").replace(" ", "")
    )


def _split_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _stem(name: str) -> str:
    raw = Path(name.replace("\\", "/")).name
    stem = Path(raw).stem or "document"
    return re.sub(r'[<>:"/\\|?*]', "_", stem)[:80] or "document"
